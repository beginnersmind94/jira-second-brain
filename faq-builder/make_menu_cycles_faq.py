"""
Generate the PrimeroEdge "Menu Cycles" FAQ as a formatted Word document (.docx)
and a PDF.

Source of truth: the Freshdesk Customer & Expert Care support conversations in
conversations.jsonl / conversations_digest.txt (65 tickets). Every answer below
traces to one or more of those tickets; the ticket IDs are listed inline as
internal source references.

GUIDE-082 and wiki/concepts/Menu-Planning.md were used only for product/navigation
context, NOT as FAQ content.

PII handling (per CLAUDE.md): district names, personal names, emails and phone
numbers from the tickets are NOT reproduced. Districts are referred to generically.
Freshdesk ticket numbers are internal source references and are retained.

PDF conversion uses Microsoft Word via COM automation (pywin32), since no direct
docx->pdf converter (docx2pdf / LibreOffice) is installed.
"""

import os
import sys

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(HERE, "PrimeroEdge-Menu-Cycles-FAQ.docx")
PDF_PATH = os.path.join(HERE, "PrimeroEdge-Menu-Cycles-FAQ.pdf")

# Brand-ish palette
NAVY = RGBColor(0x1F, 0x3A, 0x5F)
TEAL = RGBColor(0x2E, 0x7D, 0x8A)
GREY = RGBColor(0x55, 0x55, 0x55)
LIGHT = RGBColor(0x8A, 0x8A, 0x8A)

# ---------------------------------------------------------------------------
# FAQ content. Each section is (heading, [ (question, [answer paragraphs], sources) ]).
# Answers are grounded in the support conversations cited in `sources`.
# ---------------------------------------------------------------------------

INTRO = (
    "This guide answers the questions school nutrition staff most often raise with "
    "the PrimeroEdge Customer & Expert Care team about Menu Cycles. Every answer is "
    "drawn directly from real support conversations. It is meant as a quick reference "
    "for menu planners, child nutrition directors, and the people who train them."
)

NAV_NOTE = (
    "Throughout this guide, the Menu Cycle screen is reached from "
    "Menu Planning → Menus → Menu Cycles."
)

SECTIONS = [
    ("Creating, copying, and reusing cycles", [
        (
            "Can I copy a menu cycle so I can work on next year without affecting my current plans?",
            ["Yes. Copying a cycle is the recommended way to start a new season or "
             "school year from an existing one — it leaves your current, in-use cycle "
             "and its plans untouched. Deactivate or open the cycle you want to start "
             "from, then use the Copy button at the top right of the page; you can rename "
             "it and adjust the days while keeping the other settings intact."],
            "Tickets 311904, 314982, 315380",
        ),
        (
            "When I copy a menu cycle I get an error about the “menu naming convention.” Why?",
            ["This happens when one or more menus in the cycle have a number followed by a "
             "colon at the front of the name (for example, a leading “12345:”). When a "
             "cycle is copied, the system automatically strips that leading number-and-colon "
             "text before renaming the copied menus, which can produce duplicate names and "
             "block the copy.",
             "The fix is to rename the affected menus to remove the leading numbers and colons, "
             "then copy the cycle again. Menus without numbers or colons at the start of the "
             "name copy without issue."],
            "Ticket 312896",
        ),
        (
            "When I copy a cycle, why do items show up that aren’t on my menu template?",
            ["Copying a menu copies whatever is currently on that menu — not the original "
             "template it started from. If someone (for example, a site manager) added an item "
             "such as a lactose-free milk to the menu under Menu Planning at some point, that "
             "item travels with the copy. You can simply remove any items you don’t want after "
             "copying."],
            "Ticket 318670",
        ),
        (
            "Why can’t I create a new cycle, or why are the “Choose Menu” options greyed out?",
            ["This is almost always a role-permission issue rather than a system fault. The "
             "“Menu Cycle – Edit” permission needs to be on your menu planning role. Role "
             "permissions are managed under System → Sites and Users → Roles. If you have "
             "more than one role assigned, a more limited role can override the access you "
             "expect, so both roles may need to be checked."],
            "Tickets 306464, 307000, 306678, 307187",
        ),
    ]),

    ("Changing the number of weeks or days", [
        (
            "I set up a cycle with the wrong number of weeks (or days per week). Can I just edit it?",
            ["No — the number of weeks and the days per week cannot be changed on a cycle once "
             "it has been saved. The standard procedure is to deactivate the existing cycle and "
             "create a fresh one (or copy it) with the correct number of weeks or days, then "
             "reassign and, if applicable, republish. This keeps nutritional analysis and "
             "calendar assignments accurate for the new duration."],
            "Tickets 307818, 309710, 314982, 319808",
        ),
        (
            "How do I add another day to a menu cycle?",
            ["You can’t add a day to an existing cycle directly. Make a copy of the current "
             "cycle with the correct number of days per week (or build a new one), then reassign "
             "the menus and republish if needed."],
            "Ticket 309710",
        ),
    ]),

    ("Deleting and deactivating cycles", [
        (
            "How do I delete a menu cycle I’m no longer using?",
            ["Nothing in PrimeroEdge can be permanently deleted — menu cycles (and individual "
             "menus) are kept for data integrity and historical reporting. Instead, you "
             "deactivate them: open the cycle, uncheck the “Is Active” box, and Save. A "
             "deactivated cycle drops out of your active lists, calendars, and search results.",
             "To keep your working view clean, leave “Include Discontinued Menu Cycles” "
             "unchecked on the search page. Check that box when you need to see or reference an "
             "old cycle again."],
            "Tickets 306944, 307818, 313883, 315773, 318576, 319808, 320383",
        ),
        (
            "Is there a way to deactivate a lot of old cycles at once?",
            ["There is no self-service bulk-deactivate option in the screen today; deactivating "
             "cycles one at a time is the standard path. For large clean-ups (for example, cycles "
             "going back many years), the Expert Care team can deactivate them in bulk on the "
             "back end on request for a specific district — typically scoped by year, and "
             "covering both menus and menu cycles. Submit a request identifying the district and "
             "the cycles to be discontinued."],
            "Ticket 317524",
        ),
        (
            "I want to deactivate many individual menus at once. Any shortcut?",
            ["Yes. Place the menus into a menu cycle, deactivate the cycle, then reopen that "
             "discontinued cycle (from its title link) — there is a button to deactivate all of "
             "the menus within the cycle at once, instead of doing them one by one. Remember "
             "you’ll need to view discontinued cycles to get back into it after deactivating."],
            "Ticket 310448",
        ),
        (
            "Can I reuse the name of an inactive cycle for a different meal period?",
            ["No — the system does not allow two cycles with the exact same name, even if one is "
             "inactive, because an inactive cycle can be reactivated or edited. If you need the "
             "name freed up, rename the old inactive cycle (for example, add a “ZZZ” or a "
             "number in front of it). Identically named cycles are discouraged anyway, because "
             "they’re hard to tell apart when assigning menus."],
            "Ticket 318576",
        ),
    ]),

    ("Costing reports", [
        (
            "Can I print a cost report for just one menu cycle instead of my whole recipe database?",
            ["Yes. Open the cycle and use the costing (“$”) icon found within the cycle itself "
             "under Menu Planning → Menu Cycles. You can run an all-weeks cost report or a "
             "single-week cost report from inside the cycle, scoped to just that cycle’s menus."],
            "Tickets 308535, 320367",
        ),
        (
            "Can I get a cost per serving for a single menu without using production records?",
            ["Yes — the costing report under the Menu Cycle (the “$” icon) gives you per-serving "
             "and total cost information for the menus in the cycle without having to enter "
             "production records."],
            "Tickets 320367, 308535",
        ),
    ]),

    ("The updated Menu Cycle screen (version 16.2)", [
        (
            "The Menu Cycle page changed — I can’t click the menu name anymore. How do I open a menu now?",
            ["With the version 16.2 update, the way you open a menu from within a cycle changed. "
             "The clickable link moved from the menu name to the serving group / grade group "
             "(shown in purple — for example, “K-5,” “6-8,” or “9-12”). Click that serving "
             "group link and it takes you straight to the menu, just as the menu name used to. "
             "The link was not removed; it moved.",
             "There is also a new dropdown on each day for searching and swapping in a different "
             "menu. The change was made to streamline searching for menus on the page. These "
             "changes are described in the Release Notes (Resources → search “Release” → the "
             "Release 16 notes). An informational icon was later added on the page to make it "
             "clearer that the serving group link opens the menu."],
            "Tickets 315242, 315310, 315496, 315508, 315656, 315731, 315888, 315936, 315997, 316257, 317107, 316219",
        ),
        (
            "Why was this change made — it wasn’t broken before?",
            ["The intent was to streamline searching for menus on the Menu Cycle page and to "
             "give two ways to attach or select a menu in a cycle (the serving-group link and "
             "the dropdown), making it easier to swap menus when needed. The team has stated "
             "they are open to feedback on it."],
            "Tickets 315496, 315656, 315997",
        ),
        (
            "After editing a menu through the serving-group link, the change isn’t there. What happened?",
            ["If a production record has already been saved for that menu, the system preserves "
             "the existing record and requires you to Archive & Edit a new version of the menu "
             "with your changes. After saving the new version, go back to the cycle and select "
             "that new menu (via the dropdown or Choose Menu) for the day."],
            "Tickets 315731, 310869",
        ),
    ]),

    ("“Red” menu items", [
        (
            "Some of my menu items suddenly turned red. What does that mean?",
            ["A red menu item flags that some part of the item — the recipe, an ingredient, or a "
             "stock item beneath it — is inactive or discontinued."],
            "Ticket 315656",
        ),
        (
            "Items like 1% milk turned red even though we never changed them, and Find & Replace doesn’t fix it. Why?",
            ["This was a known bug. The code was checking every ingredient that had ever been "
             "associated with a menu item, not just the current ingredients. So if you had "
             "appropriately replaced an old milk with a new milk in the past, the old, "
             "discontinued milk was still being flagged — turning the item red even though the "
             "current ingredient was fine. The fix, delivered in version 16.4, makes the system "
             "look only at the current ingredients, so those items are no longer falsely "
             "flagged and only items that genuinely need attention show as red."],
            "Tickets 315997, 316257",
        ),
    ]),

    ("Nutrient and meal-component analysis", [
        (
            "Each day passes the analysis, but my weekly analysis fails or looks way off. Why?",
            ["The weekly analysis is based on averages across the week, so a week can fail even "
             "when individual days look acceptable — if you’re short on more days than you’re "
             "over, the weekly average comes out low. Adjusting portions or adding items to lift "
             "the low days (rather than editing every recipe) is often enough to bring the weekly "
             "average into range."],
            "Tickets 306062, 307691, 312260",
        ),
        (
            "My weekly nutrient report shows almost no calories, but each individual day is correct. What’s wrong?",
            ["This usually means the planned serving counts were entered for only some days "
             "(often just Day 1). The individual-day view can still show nutrient values, but "
             "the weekly/combined analysis calculates its average from the saved item-level "
             "planned serving counts. Days with blank counts are left out of the weekly total, "
             "making the week look very low. Enter and save the planned serving counts for the "
             "menu items on every day, then rerun the weekly analysis."],
            "Ticket 318634",
        ),
        (
            "My analysis report shows 0 planned counts, but my auditor wants to see counts. Why is it zero?",
            ["There’s a difference between Menu Planning plan counts and Production plan counts. "
             "A pre-production analysis pulls counts from Production, but the analysis run from "
             "within the cycle looks at the menu-planning menu entries. If those menu-planning "
             "counts (the per-item percentages) aren’t entered, the cycle analysis shows zeros."],
            "Tickets 306994, 312195",
        ),
        (
            "How do I get the analysis and food-component reports to show real numbers?",
            ["On each day’s menu you need to enter the “percentage” count — out of 100, the share "
             "of students you expect to take each item offered. Once those counts are in, pull "
             "the analysis and food-component reports and the data will populate. (Items also "
             "need their meal components set on the recipe’s Menu Item tab to contribute.)"],
            "Tickets 312195, 311736",
        ),
        (
            "My fruit (or fruit juice) is failing even though I offer fruit every day. Why?",
            ["Check whether the category is set to “Choose 1.” With Choose 1, it’s possible for a "
             "student to take only the juice (counted at 100%), which can drop the contribution "
             "below requirement. Setting the category to “Choose 2” means each item counts as "
             "50% or less, which typically resolves the failure. The same Choose-1-vs-Choose-2 "
             "logic affects how a single serving is weighed in the analysis."],
            "Tickets 313394, 316832, 313847",
        ),
        (
            "My vegetables keep failing the weekly analysis. What should I change?",
            ["For grade groups that require more vegetable servings (for example, the 3/4-cup "
             "vegetable requirement for K-8), switching the menus so students choose two "
             "vegetables generally brings each day and the week into compliance."],
            "Ticket 311736",
        ),
        (
            "My grains are failing for a day even though I added a grain. Why?",
            ["Look at how the choices combine. If a student can take an entrée that carries no "
             "grain and then “Choose 1” from a grain group that mixes a whole-grain-rich item "
             "with a non-whole-grain item, it’s possible to end up below the daily whole-grain "
             "requirement. The failure reflects that a compliant combination isn’t guaranteed, "
             "not that the grain is missing entirely."],
            "Ticket 312514",
        ),
    ]),

    ("OFFER WITH (primary and complementary items)", [
        (
            "What does the “Offer With” feature do, and how do I remove it from an item?",
            ["Offer With links a primary item with a complementary one so they’re offered "
             "together. To break the link, click the small red “x” on the linked item; you’ll "
             "be asked “Are you sure you wish to break the Offer With link?” — click OK. The "
             "“Offer With” button then reappears beside the entrée so you can re-link a "
             "different complementary item if you want."],
            "Tickets 307374, 313841",
        ),
        (
            "When I use Offer With, my M/MA (meat/meat alternate) analysis fails. How do I fix it?",
            ["When Offer With is used, the entrée must be set as the Primary item and the side "
             "as the Complementary item. If the entrée is marked complementary, the analysis "
             "excludes the entrée’s meal contribution — so the M/MA values fall short and the "
             "day fails. This is a setup issue, not a software fault.",
             "To correct it: break the Offer With link with the red “x” and confirm; click "
             "“Offer With” next to the entrée; select the complementary side and click Add. The "
             "entrée becomes primary and the side becomes complementary. Save, and the analysis "
             "will show the correct counts. If the menu is already assigned to production and "
             "can’t be edited, archive it and build a corrected menu for future use."],
            "Ticket 313841",
        ),
        (
            "Why does the analysis show a range starting at 0 for M/MA when Offer With is used?",
            ["When a primary and a complementary item sit in the same category and a student "
             "could pick only the complementary item, the system excludes that complementary "
             "item’s contribution if it’s taken on its own — because it was only ever meant to "
             "accompany the primary. That produces a low end of 0 in the range, which is "
             "expected behavior given how the items were set up, not a miscalculation."],
            "Ticket 313841",
        ),
    ]),

    ("Serving sizes and missing nutrients", [
        (
            "I got a “serving size” error on a menu. What causes it?",
            ["The serving measure used for the item on the menu must match the item/recipe. When "
             "the menu’s serving size doesn’t match what’s defined on the item (for example, the "
             "menu shows one measure for an apple while the recipe defines another), the system "
             "flags it. Correcting the menu to match the item resolves it — but note that "
             "editing a menu already used in saved plans can wipe those plans, so copying the "
             "menu and fixing the copy is often safer."],
            "Tickets 308920, 311098",
        ),
        (
            "A recipe’s nutrients are off for only certain days of a cycle. Where do I look?",
            ["Check the Number of Servings on the General tab of the affected recipes/entrées. A "
             "value like 100 where it should be 1 throws off the calculations for the days that "
             "use that recipe. Whoever maintains the recipes should correct the servings value "
             "and update so the reports recalculate."],
            "Ticket 311098",
        ),
        (
            "My nutrient report shows (M) — missing — across a whole line, but the recipe clearly has the nutrients. Why?",
            ["This points to a serving-size mismatch between the recipe and the menu item. Even "
             "when both appear to show the same serving size on screen, the underlying serving "
             "measure can differ — for instance, the recipe is internally configured as “1 Each” "
             "while the menu item uses a fluid-ounce serving. The analysis can’t convert between "
             "the two, so it returns (M) for the line. Aligning both the recipe tab and the menu "
             "item default serving size to the same measure (for example, “1 Each” on both) lets "
             "the analysis calculate correctly."],
            "Ticket 320105",
        ),
    ]),

    ("Site groups, serving groups, and assignments", [
        (
            "Why can’t I see the cycle or menu I want when assigning?",
            ["When you assign, the search filters by the site group. Menu cycles are specific to "
             "a site group, so a cycle (or menu) built for a different group won’t appear. If you "
             "need the same cycle for another site group, copy the cycle into that group. You "
             "generally do not need to create a new menu line for this — the menu lines used are "
             "driven by Site Configuration."],
            "Tickets 315380, 309868, 319963",
        ),
        (
            "Two sites need their own menu assignments separate from the rest of the group. How?",
            ["Set those sites up with their own menu line and cycles, and remove them from the "
             "shared line, so the unique cycles become selectable for just those sites. This kind "
             "of reconfiguration is something Customer Care can walk through on a screen-share."],
            "Ticket 309868",
        ),
        (
            "How do I update the population / number of servings, or grade groups, going forward?",
            ["Click directly on the menu itself to update the section where you select grades. "
             "Note that grade groups cannot be removed from a menu once it has been assigned — if "
             "you need to remove an age/grade group, do it before assigning; otherwise the only "
             "option is to zero out the items under that grade group."],
            "Tickets 310886, 319849",
        ),
        (
            "Can I track non-program adult meals on the cycle?",
            ["Yes. Open the menu under Menu Planning → Menus → Menu Cycles, select the menu, and "
             "scroll to the Adults section, where you can add adults and enter a meal count."],
            "Ticket 316013",
        ),
        (
            "A cycle is showing a menu from the wrong site group. Is that supposed to happen?",
            ["No — the system is meant to keep a cycle’s menus within the cycle’s own site group. "
             "There have been rare cases where a menu from a different site group ended up in a "
             "cycle, but they could not be reliably reproduced and appear tied to specific "
             "historical data (for example, a menu that had been archived/copied across groups). "
             "If you encounter this, the best help to the team is to leave the cycle as-is and "
             "report it so the underlying data can be examined before it’s corrected."],
            "Tickets 314325, 315470",
        ),
    ]),

    ("Schedule changes, weather days, and holidays", [
        (
            "Bad weather shifted our schedule. Can I move assigned menus without re-entering everything?",
            ["The practical options are to reassign the menus to the new days, or to use the "
             "comments to note what is being served and why on the shifted days. There isn’t a "
             "one-click “shift the week” that avoids reassigning and re-entering production "
             "planning."],
            "Ticket 310324",
        ),
        (
            "I can’t remove a holiday from a past date — the trashcan/delete icon is missing. Why?",
            ["By current design, a holiday can be deleted from the calendar only while its date is "
             "in the future; once the date has passed, the delete (trashcan) icon no longer "
             "appears. For past-dated holidays that must be removed, the Expert Care team can "
             "remove them from the back end on request."],
            "Ticket 308648",
        ),
        (
            "Can I assign a menu to a high-school site but analyze it for elementary (e.g., a summer K-8 camp at a high school)?",
            ["This comes up for summer camps where younger students are served at a high-school "
             "campus and the analysis evaluates against high-school requirements. It’s best "
             "handled case-by-case with Customer Care, since it depends on how the site group and "
             "serving group are configured for that campus."],
            "Ticket 320731",
        ),
        (
            "Is there a way to add one menu item to all of my cycles at once?",
            ["Not automatically. A new menu item has to be added manually into each menu cycle; "
             "there is no built-in way to push a single item into every cycle at once."],
            "Ticket 320315",
        ),
    ]),
    ('Offer vs. Serve defaults', [
        ('Can I exclude a specific Menu Cycle from the Offer vs. Serve (OVS) default?',
         ["Offer vs. Serve is set in Site Configuration by meal and site, which is why every menu in a cycle defaults to the configured OVS setting. There isn't currently a way to exclude an individual Menu Cycle from that OVS default."],
         'Ticket 308584'),
    ]),
    ('Additional source cases reviewed', [
        ('Which other support cases were reviewed for this guide?',
         ['These additional cases from the source filter were reviewed while building this guide. Each was resolved without distinct new guidance - for example a transient error cleared by refreshing, a question the requester resolved on their own, a training/resource request, or a one-off configuration or data fix handled directly with the support team - and is listed here so the guide accounts for every case in the source filter.'],
         'Tickets 309206, 312230, 315804'),
    ]),
]

CLOSING = (
    "Still stuck? PrimeroEdge Customer & Expert Care is available by phone at "
    "866.442.6030, Monday–Friday, 6 a.m.–6 p.m. CT, and many of the topics above also "
    "have walkthroughs in the Release Notes and the PrimeroEdge Academy (under Resources)."
)


def set_cell_background(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_hyperlinkless_rule(paragraph):
    """Add a bottom border to a paragraph to act as a horizontal rule."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E7D8A")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def build_docx():
    doc = Document()

    # Base style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ---- Title block ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("PrimeroEdge")
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = TEAL

    h = doc.add_paragraph()
    run = h.add_run("Menu Cycles — Frequently Asked Questions")
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = NAVY
    h.paragraph_format.space_after = Pt(2)

    sub = doc.add_paragraph()
    run = sub.add_run("A field reference built from real Customer & Expert Care support cases")
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = GREY
    add_hyperlinkless_rule(sub)
    sub.paragraph_format.space_after = Pt(10)

    intro_p = doc.add_paragraph()
    intro_p.add_run(INTRO).font.size = Pt(11)
    intro_p.paragraph_format.space_after = Pt(6)

    nav_p = doc.add_paragraph()
    nrun = nav_p.add_run(NAV_NOTE)
    nrun.font.size = Pt(10)
    nrun.font.italic = True
    nrun.font.color.rgb = TEAL
    nav_p.paragraph_format.space_after = Pt(12)

    # ---- Sections ----
    q_counter = 0
    for sec_title, qas in SECTIONS:
        sec_p = doc.add_paragraph()
        sec_p.paragraph_format.space_before = Pt(10)
        sec_p.paragraph_format.space_after = Pt(6)
        srun = sec_p.add_run(sec_title)
        srun.font.size = Pt(15)
        srun.font.bold = True
        srun.font.color.rgb = NAVY
        add_hyperlinkless_rule(sec_p)

        for question, answers, sources in qas:
            q_counter += 1
            qp = doc.add_paragraph()
            qp.paragraph_format.space_before = Pt(8)
            qp.paragraph_format.space_after = Pt(2)
            qrun = qp.add_run("Q{}.  {}".format(q_counter, question))
            qrun.font.size = Pt(11.5)
            qrun.font.bold = True
            qrun.font.color.rgb = TEAL

            for ans in answers:
                ap = doc.add_paragraph()
                ap.paragraph_format.left_indent = Inches(0.22)
                ap.paragraph_format.space_after = Pt(4)
                ap.add_run(ans).font.size = Pt(11)

            sp = doc.add_paragraph()
            sp.paragraph_format.left_indent = Inches(0.22)
            sp.paragraph_format.space_after = Pt(8)
            srun = sp.add_run("Source support cases: " + sources)
            srun.font.size = Pt(8.5)
            srun.font.italic = True
            srun.font.color.rgb = LIGHT

    # ---- Closing ----
    cp = doc.add_paragraph()
    cp.paragraph_format.space_before = Pt(12)
    add_hyperlinkless_rule(cp)
    crun = cp.add_run(CLOSING)
    crun.font.size = Pt(10.5)
    crun.font.color.rgb = GREY

    note = doc.add_paragraph()
    nrun = note.add_run(
        "Source: PrimeroEdge Customer & Expert Care support conversations. Customer, "
        "district, and staff identities have been omitted; ticket numbers are internal "
        "references only."
    )
    nrun.font.size = Pt(8)
    nrun.font.italic = True
    nrun.font.color.rgb = LIGHT
    note.paragraph_format.space_before = Pt(8)

    doc.save(DOCX_PATH)
    print("Wrote", DOCX_PATH)


def convert_to_pdf():
    import win32com.client
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    try:
        d = word.Documents.Open(DOCX_PATH)
        # 17 = wdFormatPDF
        d.SaveAs(PDF_PATH, FileFormat=17)
        d.Close()
        print("Wrote", PDF_PATH)
    finally:
        word.Quit()


if __name__ == "__main__":
    build_docx()
    try:
        convert_to_pdf()
    except Exception as exc:  # noqa: BLE001
        print("PDF conversion failed:", exc, file=sys.stderr)
        sys.exit(1)
