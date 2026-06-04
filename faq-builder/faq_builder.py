"""
Shared builder for PrimeroEdge function FAQs (.docx + .pdf).

Each function supplies a content spec (a dict). build(spec) writes the Word
document and converts it to PDF via Microsoft Word COM automation (pywin32),
matching the styling of the Menu Cycles / Ingredients FAQs.

Spec keys:
    title         : str   e.g. "Transfer Items"
    intro         : str
    nav_note      : str
    sections      : list[ (section_title, [ (question, [answer_paras], sources) ]) ]
    closing       : str
    source_note   : str   provenance line at the bottom
    docx_path     : str   absolute output path
    pdf_path      : str   absolute output path
    subtitle      : str   optional; defaults to the standard subtitle
"""
import sys

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
TEAL = RGBColor(0x2E, 0x7D, 0x8A)
GREY = RGBColor(0x55, 0x55, 0x55)
LIGHT = RGBColor(0x8A, 0x8A, 0x8A)

DEFAULT_SUBTITLE = "A field reference built from real Customer & Expert Care support cases"


def add_hyperlinkless_rule(paragraph):
    """Add a bottom border to a paragraph to act as a horizontal rule."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E7D8A")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def build_docx(spec):
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ---- Title block ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("PrimeroEdge")
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = TEAL

    h = doc.add_paragraph()
    run = h.add_run("{} — Frequently Asked Questions".format(spec["title"]))
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = NAVY
    h.paragraph_format.space_after = Pt(2)

    sub = doc.add_paragraph()
    run = sub.add_run(spec.get("subtitle", DEFAULT_SUBTITLE))
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = GREY
    add_hyperlinkless_rule(sub)
    sub.paragraph_format.space_after = Pt(10)

    intro_p = doc.add_paragraph()
    intro_p.add_run(spec["intro"]).font.size = Pt(11)
    intro_p.paragraph_format.space_after = Pt(6)

    nav_p = doc.add_paragraph()
    nrun = nav_p.add_run(spec["nav_note"])
    nrun.font.size = Pt(10)
    nrun.font.italic = True
    nrun.font.color.rgb = TEAL
    nav_p.paragraph_format.space_after = Pt(12)

    # ---- Sections ----
    q_counter = 0
    for sec_title, qas in spec["sections"]:
        sec_p = doc.add_paragraph()
        sec_p.paragraph_format.space_before = Pt(10)
        sec_p.paragraph_format.space_after = Pt(6)
        srun = sec_p.add_run(sec_title)
        srun.font.size = Pt(15)
        srun.font.bold = True
        srun.font.color.rgb = NAVY
        add_hyperlinkless_rule(sec_p)

        for question, answers, sources in qas:
            q_counter += 1
            qp = doc.add_paragraph()
            qp.paragraph_format.space_before = Pt(8)
            qp.paragraph_format.space_after = Pt(2)
            qrun = qp.add_run("Q{}.  {}".format(q_counter, question))
            qrun.font.size = Pt(11.5)
            qrun.font.bold = True
            qrun.font.color.rgb = TEAL

            for ans in answers:
                ap = doc.add_paragraph()
                ap.paragraph_format.left_indent = Inches(0.22)
                ap.paragraph_format.space_after = Pt(4)
                ap.add_run(ans).font.size = Pt(11)

            sp = doc.add_paragraph()
            sp.paragraph_format.left_indent = Inches(0.22)
            sp.paragraph_format.space_after = Pt(8)
            srun = sp.add_run("Source support cases: " + sources)
            srun.font.size = Pt(8.5)
            srun.font.italic = True
            srun.font.color.rgb = LIGHT

    # ---- Closing ----
    cp = doc.add_paragraph()
    cp.paragraph_format.space_before = Pt(12)
    add_hyperlinkless_rule(cp)
    crun = cp.add_run(spec["closing"])
    crun.font.size = Pt(10.5)
    crun.font.color.rgb = GREY

    note = doc.add_paragraph()
    nrun = note.add_run(spec["source_note"])
    nrun.font.size = Pt(8)
    nrun.font.italic = True
    nrun.font.color.rgb = LIGHT
    note.paragraph_format.space_before = Pt(8)

    doc.save(spec["docx_path"])
    print("Wrote", spec["docx_path"])


def convert_to_pdf(docx_path, pdf_path):
    import win32com.client
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    try:
        d = word.Documents.Open(docx_path)
        d.SaveAs(pdf_path, FileFormat=17)  # 17 = wdFormatPDF
        d.Close()
        print("Wrote", pdf_path)
    finally:
        word.Quit()


def build(spec, make_pdf=True):
    build_docx(spec)
    if make_pdf:
        try:
            convert_to_pdf(spec["docx_path"], spec["pdf_path"])
        except Exception as exc:  # noqa: BLE001
            print("PDF conversion failed:", exc, file=sys.stderr)
            return False
    return True
