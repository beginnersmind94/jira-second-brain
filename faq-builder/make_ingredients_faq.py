"""
Generate the PrimeroEdge "Ingredients" FAQ as a formatted Word document (.docx)
and a PDF.

Source of truth: the Freshdesk Customer & Expert Care support conversations in
conversations_ingredients.jsonl / conversations_ingredients_digest.txt (82
tickets, all in the Menu Planning > Ingredients category). Every answer below
traces to one or more of those tickets; the ticket IDs are listed inline as
internal source references.

Product/navigation context only (NOT FAQ content): the PrimeroEdge "Menu Planning
- Ingredients - Ingredients" Quick Step guide (Cybersoft PrimeroEdge, software
version 10.0) for breadcrumbs and tab/field labels, wiki/concepts/Menu-Planning.md,
and the PrimeroEdge Jira project. The behavior described in each answer comes from
the support conversations themselves.

PII handling (per CLAUDE.md): district names, personal names, emails and phone
numbers from the tickets are NOT reproduced. Districts are referred to
generically. Freshdesk ticket numbers are internal source references and are
retained.

PDF conversion uses Microsoft Word via COM automation (pywin32), since no direct
docx->pdf converter (docx2pdf / LibreOffice) is installed.
"""

import os
import sys

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(HERE, "deliverables", "PrimeroEdge-Ingredients-FAQ.docx")
PDF_PATH = os.path.join(HERE, "deliverables", "PrimeroEdge-Ingredients-FAQ.pdf")

# Brand-ish palette
NAVY = RGBColor(0x1F, 0x3A, 0x5F)
TEAL = RGBColor(0x2E, 0x7D, 0x8A)
GREY = RGBColor(0x55, 0x55, 0x55)
LIGHT = RGBColor(0x8A, 0x8A, 0x8A)

# ---------------------------------------------------------------------------
# FAQ content. Each section is (heading, [ (question, [answer paragraphs], sources) ]).
# Answers are grounded in the support conversations cited in `sources`.
# ---------------------------------------------------------------------------

INTRO = (
    "This guide answers the questions school nutrition staff most often raise with "
    "the PrimeroEdge Customer & Expert Care team about Ingredients — and the stock "
    "items and recipes they connect to. Every answer is drawn directly from real "
    "support conversations. It is meant as a quick reference for menu planners, child "
    "nutrition directors, inventory specialists, and the people who train them. In PrimeroEdge an ingredient is the building block a recipe consumes, and its Buying Guide tab is the bridge to the stock item that Inventory and ordering use."
)

NAV_NOTE = (
    "Throughout this guide, ingredients are added and opened from "
    "Menu Planning → Ingredients → Ingredients. Each ingredient has General, Nutrients, "
    "Allergens, Buying Guide, Documents, and History tabs; the link to the product you "
    "actually order (the stock item) is made on the Buying Guide tab."
)

SECTIONS = [
    ("How ingredients, stock items, and recipes fit together", [
        (
            "What is the difference between an ingredient, a stock item, and a recipe?",
            ["These are three connected records that play different roles. A recipe and an "
             "ingredient carry serving sizes and serving-size weights (used for nutrition and "
             "menu planning). A stock item describes how the product is purchased and ordered, "
             "and it carries the broken-unit size and weight (used for costing, ordering, "
             "production, and inventory). Stock items do not have serving measures, and recipes "
             "do not have broken-unit sizes and weights.",
             "To make everything flow through to production and the grocery list, you need a "
             "recipe, a stock item, and a link between the ingredient and the stock item. That "
             "link is made on the ingredient’s Buying Guide tab."],
            "Tickets 309390, 318086",
        ),
        (
            "Should I add a new food product as an ingredient or as a recipe?",
            ["Single, ready-to-serve products (for example, a cup cereal, a bag of chips, or a "
             "carton of milk) are entered as a single-ingredient recipe — the record is both the "
             "ingredient and the recipe. Build them in order: ingredient first, then the recipe, "
             "then link the stock item on the Buying Guide tab. A multi-item dish is built as a "
             "recipe that pulls in several ingredients."],
            "Tickets 313612, 317124, 309341",
        ),
        (
            "What is the correct order to set up a brand-new item end to end?",
            ["Create the ingredient and enter at least one serving measure (for example, 1 cup = "
             "56 grams). On the Buying Guide tab, click Edit Stock Item and complete the required "
             "fields — every field highlighted in yellow is required for production — including "
             "the broken-unit size (the smallest unit you order, e.g. 1 cup) and the broken-unit "
             "weight (the weight of that one unit). Make sure the serving measure, the broken-unit "
             "description, and how you actually order the item all agree. Finally, on the recipe "
             "select a Menu Item Category, save, and enter the meal contributions (for example, 1 "
             "whole-grain serving). Only then will the item appear on production records."],
            "Tickets 313612, 318086, 309390, 304346",
        ),
        (
            "How do I edit an existing recipe, and how do I remove an ingredient from one?",
            ["Go to Menu Planning → Ingredients → Recipes, then click the recipe name to open it. "
             "To remove an ingredient, open the Ingredients tab, find the ingredient in the list "
             "of steps, click the delete (trash) icon on that row, confirm, and Save. Keep in mind "
             "that editing a recipe does not cancel plans or orders that already exist, so account "
             "for any menus and production already built from it."],
            "Tickets 315461, 304911",
        ),
    ]),

    ("Serving sizes, measures, and weights on an ingredient", [
        (
            "How do I change an ingredient’s serving measure or grams per serving?",
            ["Serving measures live on the ingredient’s General tab. In the list of serving "
             "measures, click the pencil (edit) icon under the Edit column to change the grams, then "
             "click Update. For example, if a container’s weight changed from 150 g to 170 g, edit "
             "that serving measure rather than adding a brand-new one. Editing the grams on a serving "
             "measure flows through to the nutrients calculated for that serving. (Only serving "
             "measures you added can be edited or deleted; the system-generated measures cannot.)"],
            "Tickets 297295, 302221, 305626",
        ),
        (
            "I get an error when I try to add a serving size because one is “already in there.” What do I do?",
            ["Instead of adding a duplicate, edit the existing serving measure with the pencil icon, "
             "or delete the measure you no longer want and then enter the new one. When converting "
             "between measures, remember the math is unit-based — for example, there are 12 teaspoons "
             "in a quarter cup, so if 2 teaspoons should weigh 5 grams, the quarter-cup measure needs "
             "to be set to 30 grams for the per-serving grams to come out right."],
            "Tickets 297295, 302221",
        ),
        (
            "Why can’t I delete a serving measure — it says it is being used in a recipe?",
            ["A serving measure can’t be deleted while it is used to specify nutrients or components, "
             "or while it is referenced by any recipe — including an inactive recipe that may not "
             "show in your normal recipe list. Find and clear the last use of that measure (for "
             "example, change the measure on the inactive recipe) and it can then be removed. Because "
             "an inactive recipe can be hard to find, this is a common case where Expert Care has to "
             "locate the lingering reference."],
            "Ticket 308605",
        ),
        (
            "A serving size such as 1/2 cup won’t appear when I build the menu item. Why?",
            ["Serving sizes available on the menu item come from the recipe. Set the serving size "
             "(for example, 1/2 cup) on the recipe’s Recipe tab first and save; then that serving "
             "size becomes selectable on the menu item tab."],
            "Ticket 310141",
        ),
        (
            "How do I set fruits and vegetables to default to a 1-cup serving instead of 3/4 cup?",
            ["Create a single-ingredient recipe for the item and set its serving size to 1 cup, then "
             "save. The new serving size then shows on the nutrient view and becomes selectable in "
             "the menu’s dropdown. Using single-ingredient recipes with the serving size you want is "
             "the recommended approach for fresh fruits and vegetables."],
            "Ticket 315348",
        ),
    ]),

    ("The Buying Guide tab, stock items, and prep methods", [
        (
            "How do I link a stock item to an ingredient?",
            ["Go to Menu Planning → Ingredients → Ingredients, open the ingredient, and go to the "
             "Buying Guide tab. Select the existing stock item from the drop-down (or use Create "
             "Stock Item to build a new one), then Save — a green confirmation appears and a Remove "
             "Link button becomes available if you ever need to clear the link. Repeat this for every "
             "ingredient that needs a stock item linked, so the data flows through to production (and, "
             "for central kitchens, to the items sent out). If nutrients aren’t pulling after linking, re-save "
             "the recipe and/or click the “Update Food Required” button on the recipe; as long as "
             "the nutrients are visible on both the ingredient and the recipe, they pull through "
             "once the screen refreshes."],
            "Tickets 318086, 290768",
        ),
        (
            "The Buying Guide tab is missing on my local ingredients (but it shows on shared ones). Why?",
            ["The Buying Guide tab is controlled by a setting. Under System → Management → System "
             "Settings, on the Menu Planning tab, there is a setting called “Buying Guide – Enable "
             "Tab.” It must be turned on and saved; it can take a few hours to fully appear in "
             "PrimeroEdge afterward."],
            "Ticket 317596",
        ),
        (
            "Can you add a new prep method (for example, “Dehusked” or “Cored and Cut”) for the Buying Guide tab?",
            ["Yes — prep methods used when connecting stock items to ingredients on the Buying Guide "
             "tab come from a managed list. Send the Expert Care team the prep-method names you need "
             "and they will add them to the list."],
            "Ticket 309277",
        ),
        (
            "I’m getting an error trying to upload a document to an ingredient. What’s wrong?",
            ["Document uploads on the ingredient’s Documents tab depend on a back-end file-path "
             "setting. When uploads fail with a server error, it is handled by the Expert Care "
             "team on the back end; after the fix you may need to clear your browser cache and log "
             "out and back in for uploads to work."],
            "Ticket 319087",
        ),
    ]),

    ("Broken units (broken units per whole unit)", [
        (
            "Can I change “broken units per whole unit” on a stock item myself?",
            ["No — broken-unit fields on a stock item are not editable by district users; only the "
             "PrimeroEdge team can change them, and they do so cautiously. Submit a request if a "
             "change is genuinely needed. When a district doesn’t use inventory, the team can make "
             "the change after confirming the item and any on-hand quantities or open orders."],
            "Tickets 309341, 294965",
        ),
        (
            "Why is changing broken units per whole unit discouraged?",
            ["The system converts whole units to broken units in the background to calculate quantity "
             "and price. If broken units per whole unit is changed for an item that has had any "
             "inventory activity, all historical data for that item becomes inaccurate — what was "
             "1.5 whole units could become 2 — so inventory valuations, transfers, production records, "
             "and reports become unreliable in ways that are hard to untangle. The strong "
             "recommendation is to create a new item with the correct broken-units value and let the "
             "old one run down until it is no longer used."],
            "Tickets 309341, 313237",
        ),
        (
            "I just need to fix the case quantity so my cost is right. What should I do?",
            ["If the district doesn’t track or reconcile inventory in the system, the impact of a "
             "broken-units change is limited and the team can update it so the cost reflects the "
             "current case quantity. If you do use inventory, expect to create a new item instead. "
             "Either way, the change to a stock item’s broken units is made by the support team, not "
             "in the district’s own screens."],
            "Ticket 313237",
        ),
    ]),

    ("Food Required, production records, and the grocery list", [
        (
            "Why isn’t my single-ingredient item showing under Food Required on the production plan?",
            ["Two things are usually missing. First, the stock item linked on the ingredient’s "
             "Buying Guide tab needs a broken-unit weight — without it the system can’t calculate "
             "the amount needed. Second, the recipe must have “Show in Summary” checked so the item "
             "appears in the production plan’s food required."],
            "Ticket 295733",
        ),
        (
            "Why do weights populate on some production records but not others?",
            ["On the ingredient’s Buying Guide tab there is a checkbox to have the item weight "
             "display in Production (it sits with the linked stock item’s weight/broken-unit "
             "details). It must be checked and saved for the weight to flow to production records. "
             "If one product shows its weight and a similar one doesn’t, compare the two — the one "
             "that works will have that box checked."],
            "Ticket 309869",
        ),
        (
            "My item isn’t appearing on the grocery list or in production at all. What’s missing?",
            ["The stock item is almost certainly missing its broken-unit size and weight. Those "
             "fields are what let the system convert “1 bag” (or 1 box, 1 case) into the pounds or "
             "cups the recipe calls for; without them no calculation can be made, so the item drops "
             "off the grocery list and production. Open the ingredient’s Buying Guide tab, Edit "
             "Stock Item, and fill in the broken-unit size and weight."],
            "Tickets 309390, 318454",
        ),
        (
            "The Food Required (or Technician Worksheet) is pulling the wrong stock item or quantity. How is it fixed?",
            ["Food Required pulls from the stock item linked to the recipe’s ingredients, so start "
             "by confirming the broken-unit weight is accurate. If the worksheet still shows a stale "
             "or wrong stock item, use the “Update Food Required” button on both the stock item and "
             "the recipe, then re-save the affected production plans so the corrected item flows "
             "through next time the plans are completed."],
            "Ticket 290768",
        ),
        (
            "On the recipe, what should I enter for the ingredient quantity versus the stock-item quantity?",
            ["They must correspond to each other. If a hot pack needs 4 pieces and the recipe makes "
             "80, the ingredient quantity (320 pieces) and the stock-item quantity have to agree — "
             "leaving the stock item at a much smaller number causes it to pull the wrong amount "
             "from inventory in production. The case weight on the ingredient, the broken units on "
             "the stock item, and the recipe quantity all need to be internally consistent, or the "
             "software produces confusing results."],
            "Ticket 312257",
        ),
        (
            "A bag/case ingredient is calculating wrong because the broken-unit weight equals one serving. How do I fix it?",
            ["This is a common data error: the per-serving weight (for example, one 1.5-oz patty) "
             "was entered as the broken-unit (bag) weight too. The broken unit is the bag, so its "
             "weight must be the whole bag — for example, 53 patties × 1.5 oz = 79.5 oz per bag, "
             "with 2 bags per case. Use Edit Stock Item on the Buying Guide tab to set the broken-unit "
             "weight to the bag weight; once corrected, pricing and quantities display properly."],
            "Tickets 314306, 315204",
        ),
    ]),

    ("Costs and recipe reports", [
        (
            "An ingredient is pricing wrong inside my recipe even though the case price looks right. Why?",
            ["Check the Buying Guide tab for a stale local price. A local price entered there can "
             "override the expected case price — deleting an old, incorrect local price often "
             "corrects the recipe cost immediately. It also helps to confirm the serving-measure "
             "gram weights are right so the per-unit math lands correctly."],
            "Ticket 305626",
        ),
        (
            "My recipe report converts a measurement into something impractical (e.g., 7 Tbsp becomes “3/8 cup plus 1 Tbsp”). Can I stop that?",
            ["Yes. This is controlled by a system setting. Go to System Settings → Menu Planning, "
             "find “Method to convert Ingredient Measurements,” change it from “Use Standard "
             "Measurements” to “Use Preferred Measurement,” and Save. Refresh the recipe report and "
             "the quantities will display in the practical units (for example, tablespoons) instead "
             "of awkward fractions."],
            "Ticket 305369",
        ),
        (
            "A recipe’s total weight or scaled quantity looks wrong. How do I recalculate it?",
            ["On the recipe’s Ingredients tab, use the “(Re)calculate Weights” button to refresh the "
             "nutrient analysis and the total weight shown on the General tab; you can also use "
             "“(Re)calculate all Purchase Items” after changing ingredients or steps. Then verify "
             "each ingredient’s quantity and unit of measure on the Ingredients tab, and check the "
             "Moisture/Fat Change (%) and Waste Factor fields on the General tab, since cooking and "
             "trimming losses affect the final weight."],
            "Tickets 301629, 294367",
        ),
        (
            "When I scale a recipe, one ingredient always shows the same fixed (wrong) weight. What causes it?",
            ["This usually traces to a serving-measure weight that needs more precision. Increasing "
             "the decimal precision on the small serving measure (for example, carrying the grams "
             "for the quarter-cup serving out to four decimal places) lets the scaled quantity "
             "calculate correctly instead of sticking on one value."],
            "Ticket 294367",
        ),
    ]),

    ("Red items, inactive ingredients, and reactivating", [
        (
            "Why are some ingredients or menu items showing up in red?",
            ["Red means something in the chain is inactive — the recipe itself, one of its "
             "ingredients, or the stock item the ingredient is linked to. The item you’re looking at "
             "can appear active while a linked piece beneath it is inactive, which is enough to turn "
             "it red. (For single-ingredient products like chips, the break is often at the "
             "stock-item link.)"],
            "Tickets 317124, 305896, 318917",
        ),
        (
            "How do I bring back an ingredient or recipe that was discontinued?",
            ["For a recipe, go to Menu Planning → Recipes, search by recipe code or name, check "
             "“Include Discontinued Recipes,” and reactivate it. For an ingredient, open the "
             "ingredient (you can reach it from inside the recipe), check the “Is Active” box, and "
             "Save. Reactivating the inactive piece clears the red flag."],
            "Tickets 315439, 318917, 315360",
        ),
        (
            "Why can’t I make an ingredient inactive?",
            ["The system has several safeguards on inactivating. An ingredient can’t be made "
             "inactive while it is still linked to other records — for example a stock item link, "
             "or an open/incomplete inventory transaction. Remove those links first (and cancel any "
             "incomplete inventory), then the ingredient can be inactivated. The same idea applies "
             "to stock items: an item linked to an inactive ingredient can block deactivation until "
             "the link is removed."],
            "Tickets 312982, 307347, 315399",
        ),
    ]),

    ("Shared (state data source) ingredients and bulk changes", [
        (
            "Can I edit the allergens or nutrition on a shared (state data source) ingredient?",
            ["No — only the owner of the shared data source (for example, the state agency) can "
             "modify shared ingredient information. If you need a corrected version before they "
             "update it, open the shared ingredient and use Copy on the General tab (give the copy a "
             "unique Name and Short Name), then edit your local copy."],
            "Ticket 300604",
        ),
        (
            "How do I get a batch of shared-database items discontinued (“sent to ZZ”)?",
            ["This is a back-end clean-up the Expert Care team performs on request, identified per "
             "state data source with a list of items. For ingredients, the standard steps are: add a "
             "“ZZ” prefix to the ingredient number, discontinue the ingredient, remove its "
             "verification, and unlink the stock item. For stock items, the steps are: deactivate "
             "the stock item, change its number to a “ZZ” version, and free its SKU on the Buying "
             "Guide tab for reuse. Send the list (with item numbers) and the data source to support."],
            "Tickets 304096, 305518, 308099, 320442, 304228, 305525, 306125, 306166, 306236, 306396, 306400, 306591, 306632, 308078, 308457, 308695, 309006, 309288",
        ),
        (
            "Can you remove all of our local ingredients so we can start fresh?",
            ["The Expert Care team can remove a district’s local ingredients on request. Two "
             "important caveats: ingredients that are tied to other data such as recipes can’t be "
             "removed until those links are cleared, and PrimeroEdge discourages adopting another "
             "district’s data — if a district chooses to, it arranges that directly with the other "
             "district."],
            "Tickets 314572, 315399",
        ),
        (
            "How do we get access to our state’s shared database, or our quarterly catalog updates?",
            ["Granting a newly purchased district access to a state Department of Education shared "
             "database (ingredients, recipes, menu items) is set up by the Expert Care team. "
             "Quarterly global catalog updates are also handled by support: the district’s "
             "ingredients are exported, matched against the updated catalog, and re-imported, and "
             "the district is notified when the update is available."],
            "Tickets 317589, 299668",
        ),
        (
            "We need to import a file linking our ingredients to stock items. How does that work?",
            ["Send the linking file to the Expert Care team to import. Rows left blank (no stock "
             "item) are simply not imported, which is expected; a few rows may fail to link for "
             "specific reasons, which support will report back so you can verify and follow up."],
            "Ticket 307681",
        ),
    ]),

    ("Names, sub-ingredients, and SchoolCafe / Family Hub", [
        (
            "Where do I enter the ingredient statement (the “contains…” list) so families can see it?",
            ["Put the full ingredient statement in the Sub Ingredients section on the ingredient. "
             "Entering it there keeps it attached to the ingredient and lets it sync to SchoolCafe "
             "so it’s available on menus for parents."],
            "Tickets 316753, 313596",
        ),
        (
            "Which name shows in SchoolCafe — the Ingredient Name or the Ingredient Short Name?",
            ["The Ingredient Short Name is what pulls through to SchoolCafe. (If the short name is "
             "blank or not what you expect to see publicly, that is the field to update.)"],
            "Ticket 312024",
        ),
        (
            "What are the two ways to add ingredients, and how do I add a lot at once?",
            ["You can add them one at a time at Menu Planning → Ingredients → “Add New Ingredient,” "
             "or use the import/export tool at Menu Planning → Configuration → Import/Export, where "
             "you download an Excel template, fill it in, and import the file to add ingredients in "
             "bulk."],
            "Ticket 316753",
        ),
        (
            "When families search SchoolCafe, sub-ingredients from my sub-recipes appear too. Is that intended?",
            ["Yes — by current design, the published view includes sub-ingredients from all of a "
             "recipe’s ingredients and its sub-recipes. Limiting it to only the recipe’s own "
             "ingredients (excluding sub-recipes) would require an enhancement request, which "
             "support can escalate to the product owner; there is no committed timeline for that."],
            "Ticket 313596",
        ),
    ]),

    ("Nutrition, meal contributions, and the whole-milk rule", [
        (
            "An item’s meal contribution isn’t showing on my menu. What did I miss?",
            ["Food-component contributions are entered on the ingredient’s Nutrients tab — scroll "
             "below the per-serving nutrient fields to the Contributions section and select the "
             "quantity from the component drop-downs. Because they are entered per serving, a "
             "contribution configured for one serving (for example, a 5-count serving) does not "
             "carry to a different serving; it has to be set up for the serving size the menu "
             "actually uses, or it won’t show for that menu."],
            "Ticket 295021",
        ),
        (
            "When will whole milk / 2% milk be excluded from the saturated-fat limit (the new milk rule)?",
            ["Per the support team, the milk-rule change — saturated fat from milk not counting "
             "toward the lunch saturated-fat limit when whole and 2% milk are offered — is being "
             "delivered in the 16.4 update. It applies to lunch menus, and once in place it applies "
             "retroactively to menu analysis, so reports can be re-run afterward to show corrected "
             "results. Watch the Release Notes for the version that contains the change; the update "
             "is applied for you in the normal process."],
            "Ticket 314317",
        ),
        (
            "Whole milk isn’t an option in the menu pattern dropdown yet. Is that a setup problem on our end?",
            ["No — at the time these cases were handled the milk-pattern options (Reduced Fat and "
             "Whole Milk) were not yet available because the USDA change was still being "
             "implemented. Once delivered (in 16.4), the options become selectable and apply "
             "retroactively to your menu analysis. It is not something to fix in configuration on "
             "the district side."],
            "Ticket 314317",
        ),
    ]),
]

CLOSING = (
    "Still stuck? PrimeroEdge Customer & Expert Care is available by phone at "
    "866.442.6030, Monday–Friday, 6 a.m.–6 p.m. CT. Some changes in this guide — "
    "broken-unit edits, bulk discontinuations, shared-database access, and document-upload "
    "fixes — are completed by the Expert Care team on request rather than in your own "
    "screens, so reach out with the item numbers and the data source when you need them."
)


def set_cell_background(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_hyperlinkless_rule(paragraph):
    """Add a bottom border to a paragraph to act as a horizontal rule."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E7D8A")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def build_docx():
    doc = Document()

    # Base style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ---- Title block ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("PrimeroEdge")
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = TEAL

    h = doc.add_paragraph()
    run = h.add_run("Ingredients — Frequently Asked Questions")
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = NAVY
    h.paragraph_format.space_after = Pt(2)

    sub = doc.add_paragraph()
    run = sub.add_run("A field reference built from real Customer & Expert Care support cases")
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = GREY
    add_hyperlinkless_rule(sub)
    sub.paragraph_format.space_after = Pt(10)

    intro_p = doc.add_paragraph()
    intro_p.add_run(INTRO).font.size = Pt(11)
    intro_p.paragraph_format.space_after = Pt(6)

    nav_p = doc.add_paragraph()
    nrun = nav_p.add_run(NAV_NOTE)
    nrun.font.size = Pt(10)
    nrun.font.italic = True
    nrun.font.color.rgb = TEAL
    nav_p.paragraph_format.space_after = Pt(12)

    # ---- Sections ----
    q_counter = 0
    for sec_title, qas in SECTIONS:
        sec_p = doc.add_paragraph()
        sec_p.paragraph_format.space_before = Pt(10)
        sec_p.paragraph_format.space_after = Pt(6)
        srun = sec_p.add_run(sec_title)
        srun.font.size = Pt(15)
        srun.font.bold = True
        srun.font.color.rgb = NAVY
        add_hyperlinkless_rule(sec_p)

        for question, answers, sources in qas:
            q_counter += 1
            qp = doc.add_paragraph()
            qp.paragraph_format.space_before = Pt(8)
            qp.paragraph_format.space_after = Pt(2)
            qrun = qp.add_run("Q{}.  {}".format(q_counter, question))
            qrun.font.size = Pt(11.5)
            qrun.font.bold = True
            qrun.font.color.rgb = TEAL

            for ans in answers:
                ap = doc.add_paragraph()
                ap.paragraph_format.left_indent = Inches(0.22)
                ap.paragraph_format.space_after = Pt(4)
                ap.add_run(ans).font.size = Pt(11)

            sp = doc.add_paragraph()
            sp.paragraph_format.left_indent = Inches(0.22)
            sp.paragraph_format.space_after = Pt(8)
            srun = sp.add_run("Source support cases: " + sources)
            srun.font.size = Pt(8.5)
            srun.font.italic = True
            srun.font.color.rgb = LIGHT

    # ---- Closing ----
    cp = doc.add_paragraph()
    cp.paragraph_format.space_before = Pt(12)
    add_hyperlinkless_rule(cp)
    crun = cp.add_run(CLOSING)
    crun.font.size = Pt(10.5)
    crun.font.color.rgb = GREY

    note = doc.add_paragraph()
    nrun = note.add_run(
        "Source: PrimeroEdge Customer & Expert Care support conversations (Menu Planning "
        "> Ingredients). Customer, district, and staff identities have been omitted; ticket "
        "numbers are internal references only. Product framing was cross-checked against the jira-brain wiki concepts; all behavioral guidance traces to the support cases."
    )
    nrun.font.size = Pt(8)
    nrun.font.italic = True
    nrun.font.color.rgb = LIGHT
    note.paragraph_format.space_before = Pt(8)

    doc.save(DOCX_PATH)
    print("Wrote", DOCX_PATH)


def convert_to_pdf():
    import win32com.client
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    try:
        d = word.Documents.Open(DOCX_PATH)
        # 17 = wdFormatPDF
        d.SaveAs(PDF_PATH, FileFormat=17)
        d.Close()
        print("Wrote", PDF_PATH)
    finally:
        word.Quit()


if __name__ == "__main__":
    build_docx()
    try:
        convert_to_pdf()
    except Exception as exc:  # noqa: BLE001
        print("PDF conversion failed:", exc, file=sys.stderr)
        sys.exit(1)
